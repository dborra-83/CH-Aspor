import io
import os
from datetime import datetime
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import boto3
from weasyprint import HTML, CSS
import re

class ReportGenerator:
    """Generate DOCX and PDF reports from processed data"""
    
    def __init__(self, bucket_name: str, s3_client=None):
        self.bucket_name = bucket_name
        self.s3_client = s3_client or boto3.client('s3')
    
    def generate_docx(self, data: Dict[str, Any], output_key: str, model_type: str) -> str:
        """Generate DOCX report and upload to S3"""
        try:
            doc = Document()
            
            # Add title page
            self._add_docx_title_page(doc, data, model_type)
            
            # Add content based on model type
            if model_type == 'A':
                self._add_contragarantias_content(doc, data)
            else:
                self._add_informes_sociales_content(doc, data)
            
            # Save to memory buffer
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            
            # Upload to S3
            self.s3_client.put_object(
                Bucket=self.bucket_name,
                Key=output_key,
                Body=docx_buffer.getvalue(),
                ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
            return output_key
            
        except Exception as e:
            print(f"Error generating DOCX: {str(e)}")
            raise
    
    def generate_pdf(self, data: Dict[str, Any], output_key: str, model_type: str) -> str:
        """Generate PDF report and upload to S3"""
        try:
            # Create HTML content
            html_content = self._generate_html(data, model_type)
            
            # CSS for styling
            css = CSS(string='''
                @page {
                    size: A4;
                    margin: 2cm;
                }
                body {
                    font-family: Arial, sans-serif;
                    line-height: 1.6;
                    color: #333;
                }
                h1 {
                    color: #2c3e50;
                    border-bottom: 2px solid #3498db;
                    padding-bottom: 10px;
                }
                h2 {
                    color: #34495e;
                    margin-top: 20px;
                }
                table {
                    width: 100%;
                    border-collapse: collapse;
                    margin: 15px 0;
                }
                th, td {
                    border: 1px solid #ddd;
                    padding: 8px;
                    text-align: left;
                }
                th {
                    background-color: #f2f2f2;
                    font-weight: bold;
                }
                .alert {
                    background-color: #fff3cd;
                    border: 1px solid #ffc107;
                    padding: 10px;
                    margin: 10px 0;
                    border-radius: 5px;
                }
                .success {
                    color: green;
                }
                .error {
                    color: red;
                }
            ''')
            
            # Generate PDF
            pdf_buffer = io.BytesIO()
            HTML(string=html_content).write_pdf(pdf_buffer, stylesheets=[css])
            pdf_buffer.seek(0)
            
            # Upload to S3
            self.s3_client.put_object(
                Bucket=self.bucket_name,
                Key=output_key,
                Body=pdf_buffer.getvalue(),
                ContentType='application/pdf'
            )
            
            return output_key
            
        except Exception as e:
            print(f"Error generating PDF: {str(e)}")
            raise
    
    def _add_docx_title_page(self, doc: Document, data: Dict[str, Any], model_type: str):
        """Add title page to DOCX"""
        # Title
        title = doc.add_heading('INFORME DE EXTRACCIÓN ASPOR', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Model type
        model_name = 'CONTRAGARANTÍAS / ASPOR' if model_type == 'A' else 'INFORMES SOCIALES'
        p = doc.add_paragraph(f'Modelo: {model_name}')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date and run ID
        p = doc.add_paragraph(f'Fecha: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if 'runId' in data:
            p = doc.add_paragraph(f'ID Ejecución: {data.get("runId", "N/A")}')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
    
    def _add_contragarantias_content(self, doc: Document, data: Dict[str, Any]):
        """Add contragarantías specific content to DOCX"""
        structured = data.get('structuredData', {})
        
        # Información Societaria
        doc.add_heading('INFORMACIÓN SOCIETARIA', 1)
        content = structured.get('informacionSocietaria', 'No disponible')
        self._add_formatted_content(doc, content)
        
        # Fechas Legales Críticas
        doc.add_heading('FECHAS LEGALES CRÍTICAS Y TRAZABILIDAD NOTARIAL', 1)
        content = structured.get('fechasLegales', 'No disponible')
        self._add_formatted_content(doc, content)
        
        # Validación de Contragarantías
        doc.add_heading('VALIDACIÓN DE CONTRAGARANTÍAS', 1)
        content = structured.get('validacionContragarantias', 'No disponible')
        self._add_formatted_content(doc, content)
        
        # Apoderados por Clases
        doc.add_heading('APODERADOS POR CLASES', 1)
        content = structured.get('apoderados', 'No disponible')
        self._add_formatted_content(doc, content)
        
        # Grupos de Actuación
        doc.add_heading('GRUPOS DE ACTUACIÓN VÁLIDOS', 1)
        content = structured.get('gruposActuacion', 'No disponible')
        self._add_formatted_content(doc, content)
        
        # Alertas y Limitaciones
        doc.add_heading('ALERTAS Y LIMITACIONES', 1)
        content = structured.get('alertas', 'No disponible')
        self._add_formatted_content(doc, content)
        
        # Recomendaciones
        doc.add_heading('RECOMENDACIONES Y ACCIONES INMEDIATAS', 1)
        content = structured.get('recomendaciones', 'No disponible')
        self._add_formatted_content(doc, content)
    
    def _add_informes_sociales_content(self, doc: Document, data: Dict[str, Any]):
        """Add informes sociales specific content to DOCX"""
        structured = data.get('structuredData', {})
        
        # Antecedentes del Cliente
        doc.add_heading('ANTECEDENTES DEL CLIENTE', 1)
        content = structured.get('antecedentesCliente', 'No disponible')
        self._add_formatted_content(doc, content)
        
        # Objeto Social
        doc.add_heading('OBJETO SOCIAL', 1)
        content = structured.get('objetoSocial', 'No disponible')
        self._add_formatted_content(doc, content)
        
        # Capital Social
        doc.add_heading('CAPITAL SOCIAL', 1)
        content = structured.get('capitalSocial', 'No disponible')
        self._add_formatted_content(doc, content)
        
        # Socios o Accionistas
        doc.add_heading('SOCIOS O ACCIONISTAS Y PARTICIPACIÓN SOCIAL', 1)
        content = structured.get('socios', 'No disponible')
        self._add_formatted_content(doc, content)
        
        # Administración
        doc.add_heading('ADMINISTRACIÓN', 1)
        content = structured.get('administracion', 'No disponible')
        self._add_formatted_content(doc, content)
        
        # Directorio
        if structured.get('directorio'):
            doc.add_heading('DIRECTORIO', 1)
            content = structured.get('directorio', 'No disponible')
            self._add_formatted_content(doc, content)
        
        # Antecedentes Legales
        doc.add_heading('ANTECEDENTES LEGALES', 1)
        content = structured.get('antecedentesLegales', 'No disponible')
        self._add_formatted_content(doc, content)
        
        # Apoderados
        if structured.get('apoderados'):
            doc.add_heading('APODERADOS Y FACULTADES', 1)
            content = structured.get('apoderados', 'No disponible')
            self._add_formatted_content(doc, content)
    
    def _add_formatted_content(self, doc: Document, content: str):
        """Add formatted content to document, preserving structure"""
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check for headers
            if line.startswith('###'):
                doc.add_heading(line.replace('#', '').strip(), 3)
            elif line.startswith('##'):
                doc.add_heading(line.replace('#', '').strip(), 2)
            elif line.startswith('#'):
                doc.add_heading(line.replace('#', '').strip(), 1)
            # Check for bullet points
            elif line.startswith('-') or line.startswith('•'):
                doc.add_paragraph(line[1:].strip(), style='List Bullet')
            # Check for numbered lists
            elif re.match(r'^\d+\.', line):
                doc.add_paragraph(line, style='List Number')
            # Regular paragraph
            else:
                p = doc.add_paragraph(line)
                
                # Check for special formatting
                if '✅' in line:
                    p.add_run(' ').font.color.rgb = RGBColor(0, 128, 0)
                elif '❌' in line:
                    p.add_run(' ').font.color.rgb = RGBColor(255, 0, 0)
                elif '⚠️' in line or 'ALERTA' in line.upper():
                    p.add_run(' ').font.color.rgb = RGBColor(255, 165, 0)
    
    def _generate_html(self, data: Dict[str, Any], model_type: str) -> str:
        """Generate HTML content for PDF"""
        model_name = 'CONTRAGARANTÍAS / ASPOR' if model_type == 'A' else 'INFORMES SOCIALES'
        
        html = f'''
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Informe ASPOR - {model_name}</title>
        </head>
        <body>
            <h1>INFORME DE EXTRACCIÓN ASPOR</h1>
            <p><strong>Modelo:</strong> {model_name}</p>
            <p><strong>Fecha:</strong> {datetime.now().strftime("%d/%m/%Y %H:%M")}</p>
        '''
        
        # Add content based on model type
        structured = data.get('structuredData', {})
        
        if model_type == 'A':
            sections = [
                ('INFORMACIÓN SOCIETARIA', 'informacionSocietaria'),
                ('FECHAS LEGALES CRÍTICAS', 'fechasLegales'),
                ('VALIDACIÓN DE CONTRAGARANTÍAS', 'validacionContragarantias'),
                ('APODERADOS POR CLASES', 'apoderados'),
                ('GRUPOS DE ACTUACIÓN', 'gruposActuacion'),
                ('ALERTAS Y LIMITACIONES', 'alertas'),
                ('RECOMENDACIONES', 'recomendaciones')
            ]
        else:
            sections = [
                ('ANTECEDENTES DEL CLIENTE', 'antecedentesCliente'),
                ('OBJETO SOCIAL', 'objetoSocial'),
                ('CAPITAL SOCIAL', 'capitalSocial'),
                ('SOCIOS O ACCIONISTAS', 'socios'),
                ('ADMINISTRACIÓN', 'administracion'),
                ('ANTECEDENTES LEGALES', 'antecedentesLegales'),
                ('APODERADOS', 'apoderados')
            ]
        
        for title, key in sections:
            content = structured.get(key, 'Información no disponible')
            html += f'<h2>{title}</h2>'
            html += self._format_html_content(content)
        
        html += '''
        </body>
        </html>
        '''
        
        return html
    
    def _format_html_content(self, content: str) -> str:
        """Format content for HTML output"""
        # Convert markdown-like formatting to HTML
        content = content.replace('✅', '<span class="success">✓</span>')
        content = content.replace('❌', '<span class="error">✗</span>')
        content = content.replace('⚠️', '<span class="alert">⚠</span>')
        
        # Convert line breaks
        paragraphs = content.split('\n\n')
        formatted = ''
        
        for para in paragraphs:
            if para.strip():
                # Check if it's a list
                if para.strip().startswith('-') or para.strip().startswith('•'):
                    items = para.split('\n')
                    formatted += '<ul>'
                    for item in items:
                        if item.strip():
                            formatted += f'<li>{item.strip()[1:].strip()}</li>'
                    formatted += '</ul>'
                else:
                    formatted += f'<p>{para.strip()}</p>'
        
        return formatted