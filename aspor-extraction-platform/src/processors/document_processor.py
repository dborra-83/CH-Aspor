import boto3
import io
import PyPDF2
from docx import Document
import re
from typing import Dict, List

class DocumentProcessor:
    """Process PDF and DOCX files to extract text"""
    
    def __init__(self, bucket_name: str, s3_client=None):
        self.bucket_name = bucket_name
        self.s3_client = s3_client or boto3.client('s3')
        self.textract_client = boto3.client('textract')
    
    def extract_text(self, s3_key: str) -> str:
        """Extract text from a document in S3"""
        try:
            # Get file from S3
            response = self.s3_client.get_object(Bucket=self.bucket_name, Key=s3_key)
            file_content = response['Body'].read()
            
            # Determine file type
            if s3_key.lower().endswith('.pdf'):
                return self._extract_pdf_text(file_content, s3_key)
            elif s3_key.lower().endswith('.docx'):
                return self._extract_docx_text(file_content)
            else:
                raise ValueError(f"Unsupported file type for {s3_key}")
                
        except Exception as e:
            print(f"Error extracting text from {s3_key}: {str(e)}")
            raise
    
    def _extract_pdf_text(self, file_content: bytes, s3_key: str) -> str:
        """Extract text from PDF using PyPDF2 or Textract for scanned PDFs"""
        try:
            # Try PyPDF2 first for digital PDFs
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            # If no text extracted, likely a scanned PDF - use Textract
            if len(text.strip()) < 100:
                print(f"PDF appears to be scanned, using Textract for {s3_key}")
                return self._extract_with_textract(s3_key)
            
            return self._clean_text(text)
            
        except Exception as e:
            print(f"Error with PyPDF2, falling back to Textract: {str(e)}")
            return self._extract_with_textract(s3_key)
    
    def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(io.BytesIO(file_content))
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\t"
                    text += "\n"
            
            return self._clean_text(text)
            
        except Exception as e:
            print(f"Error extracting DOCX text: {str(e)}")
            raise
    
    def _extract_with_textract(self, s3_key: str) -> str:
        """Use AWS Textract for OCR on scanned documents"""
        try:
            # Start document text detection
            response = self.textract_client.detect_document_text(
                Document={
                    'S3Object': {
                        'Bucket': self.bucket_name,
                        'Name': s3_key
                    }
                }
            )
            
            # Extract text from blocks
            text = ""
            for block in response['Blocks']:
                if block['BlockType'] == 'LINE':
                    text += block.get('Text', '') + "\n"
            
            return self._clean_text(text)
            
        except Exception as e:
            print(f"Error with Textract: {str(e)}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Normalize encoding and remove special characters
        text = text.encode('utf-8', 'ignore').decode('utf-8')
        
        # Remove page numbers and headers/footers (common patterns)
        text = re.sub(r'Page \d+ of \d+', '', text)
        text = re.sub(r'\n\d+\n', '\n', text)
        
        return text.strip()
    
    def combine_documents(self, documents: List[Dict[str, str]]) -> str:
        """Combine multiple document texts with clear separation"""
        combined = ""
        
        for idx, doc in enumerate(documents, 1):
            combined += f"\n\n{'='*50}\n"
            combined += f"DOCUMENTO {idx}: {doc['file']}\n"
            combined += f"{'='*50}\n\n"
            combined += doc['text']
        
        return combined